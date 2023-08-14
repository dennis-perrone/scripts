#!/usr/bin/python3

# Author: Dennis Perrone
# Github: https://www.github.com/dennis-perrone
# Purpose: Create skeleton directory for colleg classes.

import os
import sys
import getpass
from pathlib import Path
from datetime import date
# #import platform
import docx

weeks = int(10)
current_user = getpass.getuser()
home = str(Path.home())
date = date.today().isoformat()

#print (platform.system())

print ('-'*60)
print ("Class Directory Creation Script")
print ('-'*60)
semester = input("What semester is this? ex. Fall 2022 would be F2022: ").upper()
class_name = input("Which class is this for? ex. IT510: ").upper()

school_dir = f'{home}/documents/college'
class_dir = f'{school_dir}/classes/{semester.lower()}-{class_name.lower()}'

def discussion_template():
    # References:
    # https://stackoverflow.com/questions/59861093/how-to-decrease-space-between-paragraphs-in-python-docx
    # https://stackoverflow.com/questions/50753314/how-to-change-font-styles-in-python-docx-for-a-sentence
    # https://stackoverflow.com/questions/63831761/how-to-change-style-of-a-title-with-python-docx
    # https://python-docx.readthedocs.io/en/latest/index.html
    discussions = int(input ("How many weekly discussions are required?: "))
    responses = int(input ("How many responses are required per discussion?: "))
    print ()
    
    doc = docx.Document()

    styles = doc.styles
    # Set color and font size of the Heading 1 style
    styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 1'].font.size = docx.shared.Pt(13)
    # Set color and font size of the Heading 2 style
    styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Heading 2'].font.size = docx.shared.Pt(12)
    # Set color and font size of the Normal style
    styles['Normal'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
    styles['Normal'].font.size = docx.shared.Pt(11)
    
    #doc = docx.Document()
    main_header = doc.add_heading(f'{class_name} Week 00 Introduction - Initial Post', level=1)
    main_header.paragraph_format.space_before = docx.shared.Pt(3)
    prompt = doc.add_heading('Discussion Prompt:', level=2)
    prompt.paragraph_format.space_before = docx.shared.Pt(3)
    doc_para = doc.add_paragraph('\n')
    resp = doc.add_heading('Discussion Response:', level=2)
    resp.paragraph_format.space_before = docx.shared.Pt(3)
    doc_para = doc.add_paragraph('\n')
    doc.save(f'{class_dir}/02-discussions/week-00-introduction.docx')

    for w in range(1, weeks +1):
        for d in range(1, discussions + 1):
            doc = docx.Document()
            styles = doc.styles
            # Set color and font size of the Heading 1 style
            styles['Heading 1'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            styles['Heading 1'].font.size = docx.shared.Pt(13)
            # Set color and font size of the Heading 2 style
            styles['Heading 2'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            styles['Heading 2'].font.size = docx.shared.Pt(12)
            # Set color and font size of the Normal style
            styles['Normal'].font.color.rgb = docx.shared.RGBColor(0, 0, 0)
            styles['Normal'].font.size = docx.shared.Pt(11)
            main_header = doc.add_heading(f'{class_name} Week {str(w).zfill(2)} Discussion {str(d).zfill(2)} - Initial Post', level=1)
            main_header.paragraph_format.space_before = docx.shared.Pt(3)
            prompt = doc.add_heading('Discussion Prompt:', level=2)
            prompt.paragraph_format.space_before = docx.shared.Pt(3)
            doc_para = doc.add_paragraph('\n')
            resp = doc.add_heading('Discussion Response:', level=2)
            resp.paragraph_format.space_before = docx.shared.Pt(3)
            doc_para = doc.add_paragraph('\n')
            for r in range(1, responses + 1):
                resp_header = doc.add_heading(f'{class_name} Week {str(w).zfill(2)} Discussion {str(d).zfill(2)} - Response Post {str(r).zfill(2)}', level=1)
                resp_header.paragraph_format.space_before = docx.shared.Pt(3)
                resp = doc.add_heading('Discussion Response:', level=2)
                resp.paragraph_format.space_before = docx.shared.Pt(3)
                doc_para = doc.add_paragraph('\n')
            doc.save(f'{class_dir}/02-discussions/week-' + str(w).zfill(2) + '-discussion-' + str(d).zfill(2) +'.docx')
    
    print(f'All discussion templates have been created for {class_name}.')

def directory_structure():
    skel_dirs = ['01-syllabus', '02-discussions', '03-notes', '04-assignments']

    if not os.path.exists(class_dir):
        os.makedirs(class_dir)
        print ()
        #print (f'{class_name} directory has been created')
        print (f'All directories have been created for {class_name}.')
        print ()
    else:
        print (f'Directory already exists.')
        exit()

    for i in skel_dirs:
        os.makedirs(os.path.join(class_dir, i))

def main():
    directory_structure()
    discussion_template()

if __name__ == "__main__":
    main()